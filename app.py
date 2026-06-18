#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
个人案件管理系统 - Flask 后端主应用
本地部署，浏览器访问，数据不上云
"""

from flask import Flask, render_template, request, jsonify, redirect, url_for, send_from_directory, send_file
from models import init_db, get_db
from datetime import datetime, date
import os
import re
import sys
import json
import uuid
import subprocess
import zipfile
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PyPDF2 import PdfReader
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False
app.config['MAX_CONTENT_LENGTH'] = None  # 不限制上传文件大小

# 支持打包后的路径解析
def get_base_path():
    """获取应用基础路径（兼容 PyInstaller 打包和源码运行）"""
    if getattr(sys, 'frozen', False):
        # PyInstaller 打包后，从 _MEIPASS 获取资源路径
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))

def get_data_path():
    """获取数据存储路径（打包后放在可执行文件旁边）"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

# 覆盖 Flask 的模板和静态文件路径
app.template_folder = os.path.join(get_base_path(), 'templates')
app.static_folder = os.path.join(get_base_path(), 'static')

# 上传目录（放在数据目录下）
UPLOAD_FOLDER = os.path.join(get_data_path(), 'data', 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 初始化数据库
init_db()


def query_db(query, args=(), one=False):
    """执行查询并返回结果"""
    db = get_db()
    cur = db.execute(query, args)
    rv = cur.fetchall()
    db.commit()
    return rv[0] if rv and one else rv


def execute_db(query, args=()):
    """执行写操作，返回lastrowid"""
    db = get_db()
    cur = db.execute(query, args)
    db.commit()
    return cur.lastrowid


# ===================== 页面路由 =====================

@app.route('/')
def index():
    """仪表盘首页"""
    return render_template('index.html')


@app.route('/cases')
def cases_page():
    return render_template('cases.html')


@app.route('/cases/<int:case_id>')
def case_detail_page(case_id):
    return render_template('case_detail.html', case_id=case_id)


@app.route('/clients')
def clients_page():
    return render_template('clients.html')


@app.route('/todos')
def todos_page():
    return render_template('todos.html')


@app.route('/templates')
def templates_page():
    return render_template('templates.html')


@app.route('/statistics')
def statistics_page():
    return render_template('statistics.html')


# ===================== 文档上传 API =====================

@app.route('/api/cases/<int:case_id>/documents', methods=['GET'])
def get_case_documents(case_id):
    """获取案件文档列表"""
    doc_type = request.args.get('doc_type', '')
    if doc_type:
        docs = query_db(
            "SELECT * FROM case_documents WHERE case_id = ? AND doc_type = ? ORDER BY created_at DESC",
            (case_id, doc_type)
        )
    else:
        docs = query_db(
            "SELECT * FROM case_documents WHERE case_id = ? ORDER BY created_at DESC",
            (case_id,)
        )
    return jsonify([dict(d) for d in docs])


@app.route('/api/cases/<int:case_id>/documents/upload', methods=['POST'])
def upload_document(case_id):
    """上传文档（传票可附带手动填写信息）"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    file = request.files['file']
    doc_type = request.form.get('doc_type', '文档')
    category = request.form.get('category', '')
    folder_path = request.form.get('folder_path', '')

    if file.filename == '':
        return jsonify({'success': False, 'message': '未选择文件'}), 400

    summons_extensions = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp', 'pdf'}
    general_extensions = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp', 'pdf',
                         'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx',
                         'txt', 'rtf', 'zip', 'rar', '7z',
                         'mp4', 'avi', 'mov', 'mkv', 'wmv', 'flv', 'webm', 'm4v', '3gp'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''

    if doc_type == '传票':
        if ext not in summons_extensions:
            return jsonify({'success': False, 'message': f'传票仅支持图片/PDF格式，不支持: {ext}'}), 400
    else:
        if ext not in general_extensions:
            return jsonify({'success': False, 'message': f'不支持的文件格式: {ext}'}), 400

    # 生成唯一文件名
    unique_name = f"{uuid.uuid4().hex}.{ext}"
    case_upload_dir = os.path.join(UPLOAD_FOLDER, str(case_id))
    os.makedirs(case_upload_dir, exist_ok=True)
    file_path = os.path.join(case_upload_dir, unique_name)
    file.save(file_path)

    # 获取文件大小
    file_size = os.path.getsize(file_path)

    # 传票手动填写字段
    summons_fields = {}
    summons_cols = [
        'hearing_date', 'hearing_time', 'hearing_location', 'hearing_court',
        'hearing_judge', 'hearing_case_number', 'case_cause', 'summons_cause',
        'summoned_party', 'summoned_address', 'contact_phone', 'clerk'
    ]
    if doc_type == '传票':
        for col in summons_cols:
            val = request.form.get(col, '').strip()
            if val:
                summons_fields[col] = val

    if summons_fields:
        columns = ['case_id', 'doc_type', 'category', 'file_name', 'file_path', 'file_size', 'folder_path'] + list(summons_fields.keys())
        values = [case_id, doc_type, category, file.filename, file_path, file_size, folder_path] + list(summons_fields.values())
        placeholders = ', '.join(['?' for _ in columns])
        execute_db(f"INSERT INTO case_documents ({', '.join(columns)}) VALUES ({placeholders})", values)
    else:
        execute_db("""
            INSERT INTO case_documents
            (case_id, doc_type, category, file_name, file_path, file_size, folder_path)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (case_id, doc_type, category, file.filename, file_path, file_size, folder_path))

    # 获取刚插入的文档ID
    doc = query_db("SELECT * FROM case_documents WHERE case_id = ? ORDER BY id DESC LIMIT 1",
                   (case_id,), one=True)

    return jsonify({
        'success': True,
        'message': '传票上传成功' if doc_type == '传票' else '文档上传成功',
        'document': dict(doc) if doc else None,
    })


@app.route('/api/cases/<int:case_id>/documents/<int:doc_id>', methods=['PUT'])
def update_document(case_id, doc_id):
    """更新文档信息（含传票手动填写字段）"""
    data = request.get_json()
    fields = []
    params = []
    editable_fields = [
        'category', 'folder_path',
        'hearing_date', 'hearing_time', 'hearing_location', 'hearing_court',
        'hearing_judge', 'hearing_case_number', 'case_cause', 'summons_cause',
        'summoned_party', 'summoned_address', 'contact_phone', 'clerk'
    ]
    for key in editable_fields:
        if key in data:
            val = data[key]
            if val is not None and val != '':
                fields.append(f"{key} = ?")
                params.append(val)
            else:
                # 允许清空字段
                fields.append(f"{key} = ?")
                params.append(None)
    if fields:
        params.extend([doc_id, case_id])
        execute_db(f"UPDATE case_documents SET {', '.join(fields)} WHERE id = ? AND case_id = ?", params)
    return jsonify({'success': True, 'message': '文档信息更新成功'})

@app.route('/api/cases/<int:case_id>/documents/<int:doc_id>', methods=['DELETE'])
def delete_document(case_id, doc_id):
    """删除文档，并联动删除对应的时间线事件和待办提醒"""
    try:
        doc = query_db("SELECT * FROM case_documents WHERE id = ? AND case_id = ?", (doc_id, case_id), one=True)
        if not doc:
            return jsonify({'success': False, 'message': '文档不存在'}), 404

        # 联动删除：由该文档生成的时间线事件
        del_te = execute_db("DELETE FROM timeline_events WHERE case_id = ? AND source_doc_id = ?", (case_id, doc_id))
        # 联动删除：由该文档生成的待办提醒
        del_td = execute_db("DELETE FROM todo_items WHERE case_id = ? AND source_doc_id = ?", (case_id, doc_id))
        # 删除文件
        try:
            if os.path.exists(doc['file_path']):
                os.remove(doc['file_path'])
        except Exception:
            pass
        execute_db("DELETE FROM case_documents WHERE id = ?", (doc_id,))
        return jsonify({
            'success': True,
            'message': '文档已删除，已同步清理关联的时间线和待办',
            'deleted_timeline': del_te,
            'deleted_todos': del_td
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'删除失败: {e}'}), 500

def add_to_calendar(title, hearing_date, hearing_time, location="", notes=""):
    """
    使用 AppleScript 直接在 macOS 日历中创建事件（无需用户确认弹窗）
    hearing_date: '2026-06-18'
    hearing_time: '09:00' or None
    返回: (success: bool, message: str)
    """
    import subprocess
    import tempfile

    # 解析日期时间
    year, month, day = hearing_date.split('-')
    if hearing_time and ':' in hearing_time:
        hour, minute = hearing_time.split(':')
    else:
        hour, minute = '9', '0'

    # 转义 AppleScript 字符串中的双引号和反斜杠
    title_esc = title.replace('\\', '\\\\').replace('"', '\\"')
    loc_esc = location.replace('\\', '\\\\').replace('"', '\\"')
    notes_esc = notes.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n').replace('\r', '')

    # AppleScript：遍历日历找到 "Work"，找不到就用第一个
    applescript = f'''
tell application "Calendar"
    set allCals to every calendar
    set targetCal to missing value

    -- 优先查找 Work 日历
    repeat with c in allCals
        if name of c is "Work" then
            set targetCal to c
            exit repeat
        end if
    end repeat

    -- 找不到 Work 就用第一个可用日历
    if targetCal is missing value then
        set targetCal to item 1 of allCals
    end if

    -- 构造开始日期（逐步设置属性，兼容中文 macOS）
    set eventDate to current date
    set year of eventDate to {int(year)}
    set month of eventDate to {int(month)}
    set day of eventDate to {int(day)}
    set hours of eventDate to {int(hour)}
    set minutes of eventDate to {int(minute)}
    set seconds of eventDate to 0

    -- 构造结束日期（+2小时）
    set endDate to current date
    set year of endDate to {int(year)}
    set month of endDate to {int(month)}
    set day of endDate to {int(day)}
    set hours of endDate to {int(hour) + 2}
    set minutes of endDate to {int(minute)}
    set seconds of endDate to 0

    -- 在目标日历中创建事件（不弹窗）
    tell targetCal
        make new event with properties {{summary:"{title_esc}", start date:eventDate, end date:endDate, location:"{loc_esc}", description:"{notes_esc}"}}
    end tell
end tell
'''

    try:
        tmp = tempfile.NamedTemporaryFile(
            mode='w', suffix='.scpt', delete=False, encoding='utf-8'
        )
        tmp.write(applescript)
        tmp.close()

        result = subprocess.run(
            ['osascript', tmp.name],
            capture_output=True, text=True, timeout=20
        )

        try:
            os.unlink(tmp.name)
        except:
            pass

        if result.returncode == 0:
            return True, '已自动添加到iCloud日历'
        else:
            err = result.stderr.strip() or result.stdout.strip() or f'returncode={result.returncode}'
            return False, f'日历添加失败: {err}'

    except subprocess.TimeoutExpired:
        try:
            os.unlink(tmp.name)
        except:
            pass
        return False, '日历添加超时'
    except Exception as e:
        return False, f'日历添加异常: {e}'

def _add_to_calendar_via_ics(title, hearing_date, hearing_time, location='', notes=''):
    """回退方案：生成 .ics 文件并通过 open 命令导入日历（需手动确认）"""
    import subprocess
    import tempfile
    from datetime import datetime, timezone, timedelta

    year, month, day = hearing_date.split('-')
    if hearing_time and ':' in hearing_time:
        hour, minute = hearing_time.split(':')
    else:
        hour, minute = '9', '0'

    tz_offset = 8
    local_dt = datetime(int(year), int(month), int(day), int(hour), int(minute), 0)
    utc_dt = local_dt - timedelta(hours=tz_offset)
    utc_end = utc_dt + timedelta(hours=2)
    dt_start = utc_dt.strftime('%Y%m%dT%H%M%S') + 'Z'
    dt_end = utc_end.strftime('%Y%m%dT%H%M%S') + 'Z'

    safe_time = (hearing_time or '0000').replace(':', '')
    uid = f"case-{hearing_date}-{safe_time}@law-case-manager"

    desc_escaped = notes.replace('\\', '\\\\').replace('\n', '\\n').replace(',', '\\,').replace(';', '\\;')
    loc_escaped = location.replace('\\', '\\\\').replace('\n', '\\n').replace(',', '\\,').replace(';', '\\;')
    title_escaped = title.replace('\\', '\\\\').replace('\n', '\\n').replace(',', '\\,').replace(';', '\\;')

    ics_lines = [
        "BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//LawCaseManager//CN",
        "CALSCALE:GREGORIAN", "METHOD:PUBLISH", "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}",
        f"DTSTART:{dt_start}", f"DTEND:{dt_end}",
        f"SUMMARY:{title_escaped}", f"LOCATION:{loc_escaped}",
        f"DESCRIPTION:{desc_escaped}",
        "STATUS:CONFIRMED", "SEQUENCE:0",
        "END:VEVENT", "END:VCALENDAR",
    ]
    ics_content = "\r\n".join(ics_lines) + "\r\n"

    ics_filename = f"开庭提醒_{hearing_date}_{safe_time}.ics"
    tmp_path = os.path.join(tempfile.gettempdir(), ics_filename)
    with open(tmp_path, 'w', encoding='utf-8') as f:
        f.write(ics_content)

    result = subprocess.run(
        ['open', '-a', 'Calendar', tmp_path],
        capture_output=True, text=True, timeout=10
    )

    import threading
    def delayed_delete(path):
        import time
        time.sleep(30)
        try:
            os.unlink(path)
        except:
            pass
    threading.Thread(target=delayed_delete, args=(tmp_path,), daemon=True).start()

    if result.returncode == 0:
        return True, '已添加到iCloud日历（请在日历弹窗中点击"好"确认）'
    else:
        return False, f'iCloud日历添加失败: {result.stderr}'




@app.route('/api/cases/<int:case_id>/documents/<int:doc_id>/image', methods=['GET'])
def get_document_image(case_id, doc_id):
    """获取文档图片（内嵌显示）"""
    doc = query_db("SELECT * FROM case_documents WHERE id = ? AND case_id = ?", (doc_id, case_id), one=True)
    if not doc:
        return jsonify({'error': '文档不存在'}), 404
    if not os.path.exists(doc['file_path']):
        return jsonify({'error': '文件不存在'}), 404
    directory = os.path.dirname(doc['file_path'])
    filename = os.path.basename(doc['file_path'])
    return send_from_directory(directory, filename)


@app.route('/api/cases/<int:case_id>/documents/<int:doc_id>/view', methods=['GET'])
def view_document_file(case_id, doc_id):
    """在浏览器中查看文档（支持图片和PDF内嵌查看）"""
    doc = query_db("SELECT * FROM case_documents WHERE id = ? AND case_id = ?", (doc_id, case_id), one=True)
    if not doc:
        return jsonify({'error': '文档不存在'}), 404
    if not os.path.exists(doc['file_path']):
        return jsonify({'error': '文件不存在'}), 404
    directory = os.path.dirname(doc['file_path'])
    filename = os.path.basename(doc['file_path'])
    # send_from_directory 会自动根据扩展名设置 Content-Type
    # PDF 会被设为 application/pdf，浏览器会内嵌显示
    return send_from_directory(directory, filename)


@app.route('/api/cases/<int:case_id>/documents/<int:doc_id>/download', methods=['GET'])
def download_document_file(case_id, doc_id):
    """下载文档文件（强制下载而非内嵌显示）"""
    doc = query_db("SELECT * FROM case_documents WHERE id = ? AND case_id = ?", (doc_id, case_id), one=True)
    if not doc:
        return jsonify({'error': '文档不存在'}), 404
    if not os.path.exists(doc['file_path']):
        return jsonify({'error': '文件不存在'}), 404
    directory = os.path.dirname(doc['file_path'])
    filename = os.path.basename(doc['file_path'])
    return send_from_directory(directory, filename, as_attachment=True, download_name=doc['file_name'])


@app.route('/api/cases/<int:case_id>/documents/download-folder', methods=['POST'])
def download_folder(case_id):
    """打包下载文件夹内所有文件为ZIP"""
    doc_ids = request.json.get('doc_ids', [])
    folder_name = request.json.get('folder_name', '文件夹')
    if not doc_ids:
        return jsonify({'error': '未选择文件'}), 400

    # 创建内存中的ZIP文件
    memory_zip = io.BytesIO()
    name_count = {}  # 处理重名文件

    with zipfile.ZipFile(memory_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
        for doc_id in doc_ids:
            doc = query_db("SELECT * FROM case_documents WHERE id = ? AND case_id = ?", (doc_id, case_id), one=True)
            if not doc or not os.path.exists(doc['file_path']):
                continue
            # 处理重名：同名文件加序号
            base_name = doc['file_name']
            if base_name in name_count:
                name_count[base_name] += 1
                name_parts = base_name.rsplit('.', 1)
                if len(name_parts) == 2:
                    archive_name = f"{name_parts[0]}_{name_count[base_name]}.{name_parts[1]}"
                else:
                    archive_name = f"{base_name}_{name_count[base_name]}"
            else:
                name_count[base_name] = 0
                archive_name = base_name
            zf.write(doc['file_path'], archive_name)

    memory_zip.seek(0)
    # 清理文件夹名中的特殊字符
    safe_name = re.sub(r'[^\w\u4e00-\u9fff]', '_', folder_name)
    return send_file(memory_zip, mimetype='application/zip', as_attachment=True,
                     download_name=f"{safe_name}.zip")




# ===================== 案件 API =====================

@app.route('/api/cases', methods=['GET'])
def get_cases():
    """获取案件列表，支持搜索和筛选"""
    keyword = request.args.get('keyword', '')
    status = request.args.get('status', '')
    case_type = request.args.get('type', '')

    query = """
        SELECT c.*, COALESCE(NULLIF(c.client_name, ''), cl.name) as client_name
        FROM cases c
        LEFT JOIN clients cl ON c.client_id = cl.id
        WHERE 1=1
    """
    params = []

    if keyword:
        query += " AND (c.case_number LIKE ? OR c.case_name LIKE ? OR c.opposing_party LIKE ?)"
        params.extend([f'%{keyword}%'] * 3)
    if status:
        query += " AND c.status = ?"
        params.append(status)
    if case_type:
        query += " AND c.case_type = ?"
        params.append(case_type)

    query += " ORDER BY c.updated_at DESC"
    cases = query_db(query, params)
    return jsonify([dict(c) for c in cases])


@app.route('/api/cases', methods=['POST'])
def create_case():
    """创建新案件，自动保存客户到客户管理"""
    data = request.get_json()
    
    # 自动保存客户到客户管理
    client_id = data.get('client_id')
    client_name = data.get('client_name', '').strip()
    client_contact = data.get('client_contact', '').strip()
    
    if client_name and not client_id:
        # 检查客户是否已存在（按名称匹配）
        existing = query_db("SELECT id, phone FROM clients WHERE name = ?", (client_name,), one=True)
        if existing:
            client_id = existing['id']
            # 如果案件填了联系方式但客户表没有，更新客户表
            if client_contact and not existing['phone']:
                execute_db("UPDATE clients SET phone = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                          (client_contact, client_id))
        else:
            # 客户不存在，自动创建
            client_id = execute_db("""
                INSERT INTO clients (name, phone, type)
                VALUES (?, ?, '个人')
            """, (client_name, client_contact))
    
    execute_db("""
        INSERT INTO cases (case_number, case_name, case_type, court, judge, status,
                          client_id, client_name, client_contact, opposing_party,
                          claim_amount, fee_amount, fee_status, invoice_status, description, create_date)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        data.get('case_number'), data.get('case_name'), data.get('case_type'),
        data.get('court'), data.get('judge'), data.get('status', '进行中'),
        client_id, client_name, client_contact,
        data.get('opposing_party'),
        data.get('claim_amount'), data.get('fee_amount'), data.get('fee_status', '未收费'),
        data.get('invoice_status', '未开票'),
        data.get('description'), data.get('create_date')
    ))
    return jsonify({'success': True, 'message': '案件创建成功'})


@app.route('/api/cases/<int:case_id>', methods=['GET'])
def get_case(case_id):
    """获取案件详情"""
    case = query_db("""
        SELECT c.*, COALESCE(NULLIF(c.client_name, ''), cl.name) as client_name,
               COALESCE(NULLIF(c.client_contact, ''), cl.phone) as client_contact_display
        FROM cases c
        LEFT JOIN clients cl ON c.client_id = cl.id
        WHERE c.id = ?
    """, (case_id,), one=True)
    if case:
        return jsonify(dict(case))
    return jsonify({'error': '案件不存在'}), 404


@app.route('/api/cases/<int:case_id>', methods=['PUT'])
def update_case(case_id):
    """更新案件，自动同步客户信息"""
    data = request.get_json()
    
    # 自动同步客户信息到客户管理
    client_name = data.get('client_name', '').strip()
    client_contact = data.get('client_contact', '').strip()
    if client_name:
        existing = query_db("SELECT id, phone FROM clients WHERE name = ?", (client_name,), one=True)
        if existing:
            data['client_id'] = existing['id']
            # 如果案件填了联系方式但客户表没有，更新客户表
            if client_contact and not existing['phone']:
                execute_db("UPDATE clients SET phone = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                          (client_contact, existing['id']))
        else:
            # 客户不存在，自动创建
            data['client_id'] = execute_db("INSERT INTO clients (name, phone, type) VALUES (?, ?, '个人')",
                      (client_name, client_contact))
    
    fields = []
    params = []
    for key in ['case_number', 'case_name', 'case_type', 'court', 'judge', 'status',
                'client_id', 'client_name', 'client_contact', 'opposing_party',
                'claim_amount', 'fee_amount', 'fee_status', 'invoice_status', 'description',
                'close_date']:
        if key in data:
            fields.append(f"{key} = ?")
            params.append(data[key])
    if fields:
        fields.append("updated_at = CURRENT_TIMESTAMP")
        params.append(case_id)
        execute_db(f"UPDATE cases SET {', '.join(fields)} WHERE id = ?", params)
    return jsonify({'success': True, 'message': '案件更新成功'})


@app.route('/api/cases/<int:case_id>', methods=['DELETE'])
def delete_case(case_id):
    """删除案件"""
    execute_db("DELETE FROM cases WHERE id = ?", (case_id,))
    return jsonify({'success': True, 'message': '案件已删除'})


# ===================== 客户 API =====================

@app.route('/api/clients', methods=['GET'])
def get_clients():
    """获取客户列表"""
    keyword = request.args.get('keyword', '')
    query = "SELECT * FROM clients WHERE 1=1"
    params = []
    if keyword:
        query += " AND (name LIKE ? OR phone LIKE ? OR id_number LIKE ?)"
        params.extend([f'%{keyword}%'] * 3)
    query += " ORDER BY updated_at DESC"
    clients = query_db(query, params)
    return jsonify([dict(c) for c in clients])


@app.route('/api/clients', methods=['POST'])
def create_client():
    """创建客户"""
    data = request.get_json()
    execute_db("""
        INSERT INTO clients (name, type, phone, email, address, id_number, notes)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (data.get('name'), data.get('type', '个人'), data.get('phone'),
          data.get('email'), data.get('address'), data.get('id_number'), data.get('notes')))
    return jsonify({'success': True, 'message': '客户创建成功'})


@app.route('/api/clients/<int:client_id>', methods=['GET'])
def get_client(client_id):
    """获取客户详情"""
    client = query_db("SELECT * FROM clients WHERE id = ?", (client_id,), one=True)
    if client:
        return jsonify(dict(client))
    return jsonify({'error': '客户不存在'}), 404


@app.route('/api/clients/<int:client_id>', methods=['PUT'])
def update_client(client_id):
    """更新客户"""
    data = request.get_json()
    fields = []
    params = []
    for key in ['name', 'type', 'phone', 'email', 'address', 'id_number', 'notes']:
        if key in data:
            fields.append(f"{key} = ?")
            params.append(data[key])
    if fields:
        fields.append("updated_at = CURRENT_TIMESTAMP")
        params.append(client_id)
        execute_db(f"UPDATE clients SET {', '.join(fields)} WHERE id = ?", params)
    return jsonify({'success': True, 'message': '客户更新成功'})


@app.route('/api/clients/<int:client_id>', methods=['DELETE'])
def delete_client(client_id):
    """删除客户"""
    execute_db("DELETE FROM clients WHERE id = ?", (client_id,))
    return jsonify({'success': True, 'message': '客户已删除'})


@app.route('/api/clients/<int:client_id>/cases', methods=['GET'])
def get_client_cases(client_id):
    """获取客户的关联案件"""
    cases = query_db("SELECT * FROM cases WHERE client_id = ? ORDER BY updated_at DESC", (client_id,))
    return jsonify([dict(c) for c in cases])


@app.route('/api/clients/export', methods=['GET'])
def export_clients():
    """导出客户列表为Excel文件"""
    # 查询所有客户
    keyword = request.args.get('keyword', '')
    query = "SELECT * FROM clients WHERE 1=1"
    params = []
    if keyword:
        query += " AND (name LIKE ? OR phone LIKE ? OR id_number LIKE ?)"
        params.extend([f'%{keyword}%'] * 3)
    query += " ORDER BY updated_at DESC"
    clients = query_db(query, params)

    # 创建Excel工作簿
    wb = Workbook()
    ws = wb.active
    ws.title = "客户列表"

    # 设置列宽
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 20
    ws.column_dimensions['F'].width = 30
    ws.column_dimensions['G'].width = 15
    ws.column_dimensions['H'].width = 30
    ws.column_dimensions['I'].width = 20

    # 表头样式
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")

    # 写表头
    headers = ['姓名/名称', '类型', '电话', '邮箱', '证件号', '地址', '备注', '关联案件数', '创建时间']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment

    # 写数据
    for row, client in enumerate(clients, 2):
        # 查询关联案件数
        case_count = query_db(
            "SELECT COUNT(*) as cnt FROM cases WHERE client_id = ?",
            (client['id'],), one=True
        )['cnt']

        ws.cell(row=row, column=1, value=client['name'])
        ws.cell(row=row, column=2, value=client['type'])
        ws.cell(row=row, column=3, value=client['phone'])
        ws.cell(row=row, column=4, value=client['email'])
        ws.cell(row=row, column=5, value=client['id_number'])
        ws.cell(row=row, column=6, value=client['address'])
        ws.cell(row=row, column=7, value=client['notes'])
        ws.cell(row=row, column=8, value=case_count)
        ws.cell(row=row, column=9, value=client['created_at'])

        # 设置数据对齐方式
        for col in range(1, 10):
            ws.cell(row=row, column=col).alignment = Alignment(vertical="center")

    # 生成文件名
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"客户列表_{timestamp}.xlsx"

    # 保存到内存
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


# ===================== 工作记录 API =====================

@app.route('/api/work-records', methods=['GET'])
def get_work_records():
    """获取工作记录"""
    case_id = request.args.get('case_id')
    if case_id:
        records = query_db("SELECT * FROM work_records WHERE case_id = ? ORDER BY date DESC", (case_id,))
    else:
        records = query_db("SELECT * FROM work_records ORDER BY date DESC LIMIT 100")
    return jsonify([dict(r) for r in records])


@app.route('/api/work-records', methods=['POST'])
def create_work_record():
    """创建工作记录"""
    data = request.get_json()
    execute_db("""
        INSERT INTO work_records (case_id, date, content, hours, category)
        VALUES (?, ?, ?, ?, ?)
    """, (data.get('case_id'), data.get('date'), data.get('content'),
          data.get('hours'), data.get('category')))
    return jsonify({'success': True, 'message': '工作记录添加成功'})


@app.route('/api/work-records/<int:record_id>', methods=['PUT'])
def update_work_record(record_id):
    """更新工作记录"""
    data = request.get_json()
    fields = []
    params = []
    for key in ['date', 'content', 'hours', 'category']:
        if key in data:
            fields.append(f"{key} = ?")
            params.append(data[key])
    if fields:
        params.append(record_id)
        execute_db(f"UPDATE work_records SET {', '.join(fields)} WHERE id = ?", params)
    return jsonify({'success': True, 'message': '工作记录更新成功'})


@app.route('/api/work-records/<int:record_id>', methods=['DELETE'])
def delete_work_record(record_id):
    """删除工作记录"""
    execute_db("DELETE FROM work_records WHERE id = ?", (record_id,))
    return jsonify({'success': True, 'message': '工作记录已删除'})


# ===================== 待办事项 API =====================

@app.route('/api/todos', methods=['GET'])
def get_todos():
    """获取待办事项"""
    status = request.args.get('status', '')
    case_id = request.args.get('case_id', '')
    query = """
        SELECT t.*, c.case_name
        FROM todo_items t
        LEFT JOIN cases c ON t.case_id = c.id
        WHERE 1=1
    """
    params = []
    if status:
        query += " AND t.status = ?"
        params.append(status)
    if case_id:
        query += " AND t.case_id = ?"
        params.append(int(case_id))
    query += " ORDER BY CASE t.priority WHEN '紧急' THEN 1 WHEN '重要' THEN 2 ELSE 3 END, t.deadline ASC"
    todos = query_db(query, params)
    return jsonify([dict(t) for t in todos])


@app.route('/api/todos', methods=['POST'])
def create_todo():
    """创建待办事项"""
    data = request.get_json()
    execute_db("""
        INSERT INTO todo_items (case_id, title, description, deadline, priority, status, reminder_date)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (data.get('case_id'), data.get('title'), data.get('description'),
          data.get('deadline'), data.get('priority', '普通'),
          data.get('status', '待办'), data.get('reminder_date')))
    return jsonify({'success': True, 'message': '待办事项创建成功'})


@app.route('/api/todos/<int:todo_id>', methods=['PUT'])
def update_todo(todo_id):
    """更新待办事项"""
    data = request.get_json()
    fields = []
    params = []
    for key in ['title', 'description', 'deadline', 'priority', 'status', 'reminder_date', 'case_id']:
        if key in data:
            fields.append(f"{key} = ?")
            params.append(data[key])
    if fields:
        fields.append("updated_at = CURRENT_TIMESTAMP")
        params.append(todo_id)
        execute_db(f"UPDATE todo_items SET {', '.join(fields)} WHERE id = ?", params)
    return jsonify({'success': True, 'message': '待办事项更新成功'})


@app.route('/api/todos/<int:todo_id>', methods=['DELETE'])
def delete_todo(todo_id):
    """删除待办事项"""
    execute_db("DELETE FROM todo_items WHERE id = ?", (todo_id,))
    return jsonify({'success': True, 'message': '待办事项已删除'})


# ===================== 案件跟进 API =====================

@app.route('/api/followups', methods=['GET'])
def get_followups():
    """获取跟进记录"""
    case_id = request.args.get('case_id')
    if case_id:
        followups = query_db("SELECT * FROM case_followups WHERE case_id = ? ORDER BY date DESC", (case_id,))
    else:
        followups = query_db("SELECT * FROM case_followups ORDER BY date DESC LIMIT 50")
    return jsonify([dict(f) for f in followups])


@app.route('/api/followups', methods=['POST'])
def create_followup():
    """创建跟进记录"""
    data = request.get_json()
    execute_db("""
        INSERT INTO case_followups (case_id, date, content, contact_person, contact_phone, result, next_action)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (data.get('case_id'), data.get('date'), data.get('content'),
          data.get('contact_person', ''), data.get('contact_phone', ''),
          data.get('result'), data.get('next_action')))
    return jsonify({'success': True, 'message': '跟进记录添加成功'})


@app.route('/api/followups/<int:fuid>', methods=['DELETE'])
def delete_followup(fuid):
    """删除跟进记录"""
    execute_db("DELETE FROM case_followups WHERE id = ?", (fuid,))
    return jsonify({'success': True, 'message': '跟进记录已删除'})


@app.route('/api/followups/<int:fuid>', methods=['PUT'])
def update_followup(fuid):
    """更新跟进记录"""
    data = request.get_json()
    fields = []
    params = []
    for key in ['date', 'content', 'contact_person', 'contact_phone', 'result', 'next_action']:
        if key in data:
            fields.append(f"{key} = ?")
            params.append(data[key])
    if fields:
        params.append(fuid)
        execute_db(f"UPDATE case_followups SET {', '.join(fields)} WHERE id = ?", params)
    return jsonify({'success': True, 'message': '跟进记录更新成功'})


# ===================== 文书模板 API =====================

@app.route('/api/templates', methods=['GET'])
def get_templates():
    """获取文书模板列表"""
    category = request.args.get('category', '')
    query = "SELECT * FROM doc_templates WHERE 1=1"
    params = []
    if category:
        query += " AND category = ?"
        params.append(category)
    query += " ORDER BY updated_at DESC"
    templates = query_db(query, params)
    return jsonify([dict(t) for t in templates])


@app.route('/api/templates', methods=['POST'])
def create_template():
    """创建文书模板"""
    data = request.get_json()
    file_path = data.get('file_path', '')
    execute_db("""
        INSERT INTO doc_templates (name, category, content, notes, file_path)
        VALUES (?, ?, ?, ?, ?)
    """, (data.get('name'), data.get('category'), data.get('content'), data.get('notes'), file_path))
    return jsonify({'success': True, 'message': '模板创建成功'})


@app.route('/api/templates/<int:tpl_id>', methods=['GET'])
def get_template(tpl_id):
    """获取模板详情"""
    tpl = query_db("SELECT * FROM doc_templates WHERE id = ?", (tpl_id,), one=True)
    if tpl:
        return jsonify(dict(tpl))
    return jsonify({'error': '模板不存在'}), 404


@app.route('/api/templates/<int:tpl_id>', methods=['PUT'])
def update_template(tpl_id):
    """更新模板"""
    data = request.get_json()
    fields = []
    params = []
    for key in ['name', 'category', 'content', 'notes', 'file_path']:
        if key in data:
            fields.append(f"{key} = ?")
            params.append(data[key])
    if fields:
        fields.append("updated_at = CURRENT_TIMESTAMP")
        params.append(tpl_id)
        execute_db(f"UPDATE doc_templates SET {', '.join(fields)} WHERE id = ?", params)
    return jsonify({'success': True, 'message': '模板更新成功'})


@app.route('/api/templates/<int:tpl_id>', methods=['DELETE'])
def delete_template(tpl_id):
    """删除模板"""
    execute_db("DELETE FROM doc_templates WHERE id = ?", (tpl_id,))
    return jsonify({'success': True, 'message': '模板已删除'})


@app.route('/api/templates/<int:tpl_id>/download', methods=['GET'])
def download_template(tpl_id):
    """下载模板（有PDF原文件则下载PDF，否则生成Word）"""
    tpl = query_db("SELECT * FROM doc_templates WHERE id = ?", (tpl_id,), one=True)
    if not tpl:
        return jsonify({'error': '模板不存在'}), 404

    # 如果有PDF原文件，直接返回PDF
    if tpl['file_path'] and os.path.exists(tpl['file_path']):
        safe_name = re.sub(r'[^\w\u4e00-\u9fff]', '_', tpl['name'])
        return send_file(tpl['file_path'], mimetype='application/pdf',
                         as_attachment=True, download_name=f"{safe_name}.pdf")

    doc = Document()
    # 设置默认字体为宋体（需同时设置西文和东亚字体）
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    # 设置东亚字体为宋体（关键：确保中文也使用宋体）
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    title_para = doc.add_paragraph(tpl['name'])
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.name = '宋体'
    title_run.bold = True
    # 标题也设置东亚字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 正文内容 - 逐行添加，保留空行
    content = tpl['content'] or ''
    lines = content.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(22)  # 1.5倍行距
        p.paragraph_format.space_after = Pt(0)
        # 高亮 {变量名} 占位符
        if '{' in line and '}' in line:
            remaining = line
            while '{' in remaining and '}' in remaining:
                start = remaining.index('{')
                end = remaining.index('}', start)
                if start > 0:
                    normal_run = p.add_run(remaining[:start])
                    normal_run.font.name = '宋体'
                    normal_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                var_run = p.add_run(remaining[start:end+1])
                var_run.font.highlight_color = 7  # 黄色高亮
                var_run.font.name = '宋体'
                var_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                remaining = remaining[end+1:]
            if remaining:
                normal_run = p.add_run(remaining)
                normal_run.font.name = '宋体'
                normal_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            run = p.add_run(line)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 使用说明
    if tpl['notes']:
        doc.add_paragraph('')  # 空行分隔
        notes_para = doc.add_paragraph(f'【使用说明】{tpl["notes"]}')
        notes_run = notes_para.runs[0]
        notes_run.font.size = Pt(10)
        notes_run.font.name = '宋体'
        notes_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 生成到内存
    memory_file = io.BytesIO()
    doc.save(memory_file)
    memory_file.seek(0)

    # 文件名处理
    safe_name = re.sub(r'[^\w\u4e00-\u9fff]', '_', tpl['name'])
    return send_file(memory_file, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name=f"{safe_name}.docx")


@app.route('/api/templates/<int:tpl_id>/preview-pdf', methods=['GET'])
def preview_template_pdf(tpl_id):
    """预览PDF模板（内联显示）"""
    tpl = query_db("SELECT * FROM doc_templates WHERE id = ?", (tpl_id,), one=True)
    if not tpl:
        return jsonify({'error': '模板不存在'}), 404
    if not tpl['file_path'] or not os.path.exists(tpl['file_path']):
        return jsonify({'error': '无PDF文件'}), 404
    return send_file(tpl['file_path'], mimetype='application/pdf', as_attachment=False)


@app.route('/api/templates/upload', methods=['POST'])
def upload_template():
    """上传Word或PDF文件创建模板"""
    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '未选择文件'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ('.docx', '.doc', '.pdf'):
        return jsonify({'error': '仅支持 .docx 和 .pdf 格式文件'}), 400

    if ext == '.doc':
        return jsonify({'error': '暂不支持 .doc 旧格式，请转换为 .docx 后上传'}), 400

    try:
        name = os.path.splitext(file.filename)[0]
        category = request.form.get('category', '其他')
        notes = request.form.get('notes', '')
        file_path = ''

        if ext == '.pdf':
            # 保存PDF原文件
            template_dir = os.path.join(UPLOAD_FOLDER, 'templates')
            os.makedirs(template_dir, exist_ok=True)
            safe_name = re.sub(r'[^\w\u4e00-\u9fff]', '_', name)
            temp_filename = f"{safe_name}_{uuid.uuid4().hex[:8]}{ext}"
            file_path = os.path.join(template_dir, temp_filename)
            file.seek(0)
            file.save(file_path)

            # 使用 PyPDF2 提取 PDF 文本
            reader = PdfReader(file_path)
            paragraphs = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        paragraphs.append(line.strip())
            content = '\n'.join(paragraphs)
        else:
            # docx 文件
            doc = Document(file)
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            content = '\n'.join(paragraphs)

        # 检查是否重名，重名则追加序号
        existing = query_db("SELECT id FROM doc_templates WHERE name = ?", (name,), one=True)
        if existing:
            count = query_db("SELECT COUNT(*) as cnt FROM doc_templates WHERE name LIKE ?", (name + '%',))[0]['cnt']
            name = f"{name}_{count}"

        execute_db("""
            INSERT INTO doc_templates (name, category, content, notes, file_path)
            VALUES (?, ?, ?, ?, ?)
        """, (name, category, content, notes, file_path))
        return jsonify({'success': True, 'message': '模板上传成功', 'name': name})
    except Exception as e:
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 500


@app.route('/api/templates/parse-upload', methods=['POST'])
def parse_template_upload():
    """上传Word或PDF文件解析内容（不保存，返回文本供预览；PDF会暂存文件）"""
    if 'file' not in request.files:
        return jsonify({'error': '未选择文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '未选择文件'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext == '.doc':
        return jsonify({'error': '暂不支持 .doc 旧格式，请转换为 .docx 后上传'}), 400
    if ext not in ('.docx', '.pdf'):
        return jsonify({'error': '仅支持 .docx 和 .pdf 格式文件'}), 400

    try:
        name = os.path.splitext(file.filename)[0]
        is_pdf = (ext == '.pdf')
        saved_path = ''

        if ext == '.pdf':
            # 保存PDF原文件到临时目录
            template_dir = os.path.join(UPLOAD_FOLDER, 'templates')
            os.makedirs(template_dir, exist_ok=True)
            safe_name = re.sub(r'[^\w\u4e00-\u9fff]', '_', name)
            temp_filename = f"{safe_name}_{uuid.uuid4().hex[:8]}{ext}"
            saved_path = os.path.join(template_dir, temp_filename)
            file.seek(0)
            file.save(saved_path)

            # 使用 PyPDF2 提取 PDF 文本
            reader = PdfReader(saved_path)
            paragraphs = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        paragraphs.append(line.strip())
            content = '\n'.join(paragraphs)
        else:
            # docx 文件
            doc = Document(file)
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            content = '\n'.join(paragraphs)

        return jsonify({'success': True, 'name': name, 'content': content, 'is_pdf': is_pdf, 'saved_path': saved_path})
    except Exception as e:
        return jsonify({'error': f'文件解析失败: {str(e)}'}), 500


# ===================== 时间线 API =====================

@app.route('/api/timeline', methods=['GET'])
def get_timeline():
    """获取案件时间线"""
    case_id = request.args.get('case_id')
    if case_id:
        events = query_db("""
            SELECT * FROM timeline_events WHERE case_id = ? ORDER BY date ASC
        """, (case_id,))
    else:
        events = query_db("SELECT * FROM timeline_events ORDER BY date DESC LIMIT 50")
    return jsonify([dict(e) for e in events])


@app.route('/api/timeline', methods=['POST'])
def create_timeline_event():
    """创建时间线事件"""
    data = request.get_json()
    execute_db("""
        INSERT INTO timeline_events (case_id, date, title, description, event_type)
        VALUES (?, ?, ?, ?, ?)
    """, (data.get('case_id'), data.get('date'), data.get('title'),
          data.get('description'), data.get('event_type')))
    return jsonify({'success': True, 'message': '时间线事件添加成功'})


@app.route('/api/timeline/<int:event_id>', methods=['PUT'])
def update_timeline_event(event_id):
    """更新时间线事件"""
    data = request.get_json()
    fields = []
    params = []
    for key in ['date', 'title', 'description', 'event_type']:
        if key in data:
            fields.append(f"{key} = ?")
            params.append(data[key])
    if fields:
        params.append(event_id)
        execute_db(f"UPDATE timeline_events SET {', '.join(fields)} WHERE id = ?", params)
    return jsonify({'success': True, 'message': '时间线事件更新成功'})


@app.route('/api/timeline/<int:event_id>', methods=['DELETE'])
def delete_timeline_event(event_id):
    """删除时间线事件"""
    execute_db("DELETE FROM timeline_events WHERE id = ?", (event_id,))
    return jsonify({'success': True, 'message': '时间线事件已删除'})


# ===================== 统计 API =====================

@app.route('/api/statistics/overview', methods=['GET'])
def get_statistics():
    """获取统计概览"""
    total_cases = query_db("SELECT COUNT(*) as cnt FROM cases", one=True)['cnt']
    active_cases = query_db("SELECT COUNT(*) as cnt FROM cases WHERE status IN ('一审进行中', '二审进行中', '执行进行中')", one=True)['cnt']
    closed_cases = query_db("SELECT COUNT(*) as cnt FROM cases WHERE status IN ('一审已结案', '二审已结案', '执行完毕', '已结案')", one=True)['cnt']

    pending_todos = query_db("SELECT COUNT(*) as cnt FROM todo_items WHERE status != '已完成'", one=True)['cnt']
    overdue_todos = query_db(
        "SELECT COUNT(*) as cnt FROM todo_items WHERE status != '已完成' AND deadline < date('now')",
        one=True)['cnt']

    total_claim = query_db("SELECT COALESCE(SUM(claim_amount), 0) as total FROM cases", one=True)['total']
    total_fee = query_db("SELECT COALESCE(SUM(fee_amount), 0) as total FROM cases", one=True)['total']
    unpaid_fee = query_db(
        "SELECT COALESCE(SUM(fee_amount), 0) as total FROM cases WHERE fee_status IN ('未收费', '部分收费')",
        one=True)['total']

    type_dist = query_db("""
        SELECT case_type, COUNT(*) as count
        FROM cases GROUP BY case_type ORDER BY count DESC
    """)

    status_dist = query_db("""
        SELECT status, COUNT(*) as count
        FROM cases GROUP BY status
    """)

    fee_dist = query_db("""
        SELECT fee_status, COUNT(*) as count, COALESCE(SUM(fee_amount), 0) as amount
        FROM cases GROUP BY fee_status
    """)

    monthly_hours = query_db("""
        SELECT COALESCE(SUM(hours), 0) as total
        FROM work_records
        WHERE date >= date('now', '-30 days')
    """, one=True)['total']

    upcoming_todos = query_db("""
        SELECT t.*, c.case_name
        FROM todo_items t
        LEFT JOIN cases c ON t.case_id = c.id
        WHERE t.status != '已完成'
        ORDER BY CASE WHEN t.deadline IS NULL THEN 0 ELSE 1 END, t.deadline ASC
        LIMIT 10
    """)

    return jsonify({
        'total_cases': total_cases,
        'active_cases': active_cases,
        'closed_cases': closed_cases,
        'pending_todos': pending_todos,
        'overdue_todos': overdue_todos,
        'total_claim': total_claim,
        'total_fee': total_fee,
        'unpaid_fee': unpaid_fee,
        'monthly_hours': monthly_hours,
        'type_distribution': [dict(r) for r in type_dist],
        'status_distribution': [dict(r) for r in status_dist],
        'fee_distribution': [dict(r) for r in fee_dist],
        'upcoming_todos': [dict(r) for r in upcoming_todos],
    })


@app.route('/api/case-types', methods=['GET'])
def get_case_types():
    """获取案件案由列表（从已有案件中提取去重值）"""
    types = query_db("SELECT DISTINCT case_type AS name FROM cases WHERE case_type IS NOT NULL AND case_type != '' ORDER BY case_type")
    return jsonify([dict(t) for t in types])




# ===================== 启动 =====================

if __name__ == '__main__':
    import webbrowser
    import threading

    print("\n" + "=" * 50)
    print("  个人案件管理系统")
    print("  本地部署 | 数据不上云 | 隐私优先")
    print("=" * 50)

    # 自动打开浏览器
    def open_browser():
        webbrowser.open('http://127.0.0.1:5066')
    threading.Timer(1.5, open_browser).start()

    print(f"\n  请在浏览器中访问: http://127.0.0.1:5066\n")
    app.run(host='127.0.0.1', port=5066, debug=False)
